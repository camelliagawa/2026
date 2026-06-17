"""Generate Word document explaining the polynomial fit blade curve method."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text='', bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if size:
            run.font.size = Pt(size)
    return p

def add_formula(doc, formula_text):
    """Add a formula paragraph with centered alignment and monospace font."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font and light background."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Add shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def add_table_row(table, cells_data):
    row = table.add_row()
    for i, (text, bold) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        if bold:
            run.bold = True
    return row

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Yu Mincho'
style.font.size = Pt(11)

# Set page margins
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# ============================================================
# Title
# ============================================================
title = doc.add_heading('刃渡り曲線の推定方法', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('端点拘束付き多項式最小二乗フィットによる滑らか曲線推定')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(13)
subtitle.runs[0].italic = True

doc.add_paragraph()

# ============================================================
# 1. 概要
# ============================================================
add_heading(doc, '1. 概要', level=1)

p = doc.add_paragraph()
p.add_run(
    'ロボット研磨ルートを生成するためには、包丁の刃線（アゴから切先にかけての刃渡り曲線）を'
    '正確かつ滑らかに推定する必要がある。本アプリでは、ユーザーが画像上でアゴ（刃元）と'
    '切先を手動指定した後、以下の手順で刃渡り曲線を求める。'
)

steps = [
    ('① エッジ追跡', 'エッジ検出画像から各 x 位置における刃先ピクセルを走査する'),
    ('② ギャップ補間', '走査で見つからなかった x 位置を線形補間で埋める'),
    ('③ メディアンフィルタ', 'スパイク状外れ値を除去する'),
    ('④ 多項式フィット', '端点拘束付き最小二乗法で滑らか曲線を当てはめる'),
]
for step, desc in steps:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(step + '：')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()

# ============================================================
# 2. エッジ追跡
# ============================================================
add_heading(doc, '2. エッジ追跡（ステップ①）', level=1)

p = doc.add_paragraph()
p.add_run(
    'アゴ座標 (x₀, y₀) と切先座標 (x₁, y₁) を結ぶ「弦」を基準として、'
    '各 x 位置でエッジ画像を縦方向に走査し、最も刃先側にあるエッジピクセルを検出する。'
)

add_heading(doc, '2.1 走査ウィンドウの非対称設計', level=2)

p = doc.add_paragraph()
p.add_run(
    '弦上の参照点 (x, refY) を中心に、上方向 UP ピクセル・下方向 DOWN ピクセルの'
    '非対称ウィンドウで走査する。'
)

add_formula(doc, 'UP   = max(8,  round(W × 0.006))')
add_formula(doc, 'DOWN = max(20, round(W × 0.015))')
add_formula(doc, 'W = |x₁ - x₀|（アゴ〜切先の水平距離、ピクセル）')

p = doc.add_paragraph()
p.add_run('この非対称設計の根拠：')

items = [
    ('下方向 DOWN を広くとる理由',
     '刃先は弦よりも常に刃側（下方）に位置するため、刃先ピクセルを見逃さないよう'
     '広めの探索範囲が必要。'),
    ('上方向 UP を狭くとる理由',
     '弦より上（棟側）にあるピクセルは刃線ではなく包丁本体の輪郭ノイズである可能性が高い。'
     '不要な点を拾わないよう探索を制限する。'),
    ('比率をピクセル固定ではなく W の比率にする理由',
     '画像の解像度・撮影距離によって包丁サイズが変わっても、探索範囲が適切にスケールする。'),
    ('最小値（max）を設ける理由',
     '包丁が画像内で非常に小さく写っている場合でも、最低限のウィンドウ幅を保証する。'),
]
for title_text, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_text + '：')
    run.bold = True
    p.add_run(desc)

add_heading(doc, '2.2 走査方向', level=2)

p = doc.add_paragraph()
p.add_run(
    '走査は「下から上」（y の大きい方から小さい方）の順に行い、'
    '最初に見つかったエッジピクセルを採用する。これにより、最も刃先側（最大 y 側）の'
    'エッジが選択される。'
)

add_code_block(doc,
'for y in range(refY + DOWN, refY - UP - 1, -1):   # 下→上\n'
'    if isEdge(x, y):\n'
'        found = y\n'
'        break'
)

doc.add_paragraph()

# ============================================================
# 3. ギャップ補間
# ============================================================
add_heading(doc, '3. ギャップ補間（ステップ②）', level=1)

p = doc.add_paragraph()
p.add_run(
    'エッジ画像のノイズや包丁表面の反射により、一部の x 位置でエッジピクセルが'
    '見つからない（yRaw[s] = null）場合がある。このような「ギャップ」を'
    '隣接する検出済み点を結ぶ線形補間で埋める。'
)

p = doc.add_paragraph()
p.add_run('補間の効果：')

items = [
    'ギャップ区間でも曲線が途切れずに連続する',
    '後段のメディアンフィルタや多項式フィットが全区間で適用可能になる',
    'アゴ点 (yRaw[0] = y₀) と切先点 (yRaw[xSteps] = y₁) は常に確定値として固定されるため、補間の端点が必ず存在する',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# ============================================================
# 4. メディアンフィルタ
# ============================================================
add_heading(doc, '4. メディアンフィルタ（ステップ③）', level=1)

p = doc.add_paragraph()
p.add_run(
    'エッジ追跡後の生データ yRaw には、エッジ検出の誤検知によるスパイク状の外れ値が'
    '残ることがある。これを除去するためにメディアンフィルタを適用する。'
)

add_heading(doc, '4.1 ウィンドウサイズ', level=2)

add_formula(doc, 'window = max(3, round(xSteps × 0.01))')
add_formula(doc, 'xSteps = |x₁ - x₀|（アゴ〜切先の水平距離）')

p = doc.add_paragraph()
p.add_run(
    'ウィンドウサイズを xSteps の 1% に設定することで、局所的なスパイク（ノイズ）を'
    '除去しつつ、包丁の全体的な刃線形状（数%スケールの変化）は保持する。'
    'ウィンドウが広すぎると刃先の急激な曲率変化（切先付近）を平滑化しすぎてしまうため、'
    '控えめなサイズにとどめる。'
)

add_heading(doc, '4.2 移動平均との比較', level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '手法'
hdr[1].text = '特徴'
hdr[2].text = '包丁への適合性'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('移動平均', '各点の近傍平均を取る\n局所的スムージング', '刃先の不規則なノイズを局所的に低減するが、\nスパイク1点が窓全体に影響を与える'),
    ('メディアンフィルタ', '各点の近傍中央値を取る\nスパイク耐性が高い', '単発の誤検知ピクセルを効果的に除去。\n刃線の大域形状は維持される'),
]
for row_data in data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        row.cells[i].text = text

doc.add_paragraph()

# ============================================================
# 5. 多項式フィット（メイン手法）
# ============================================================
add_heading(doc, '5. 端点拘束付き多項式最小二乗フィット（ステップ④）', level=1)

p = doc.add_paragraph()
p.add_run(
    'メディアンフィルタ後のデータ yMed に対し、アゴと切先を正確に通る滑らかな曲線を'
    '多項式最小二乗法で当てはめる。これが本手法の核心部分である。'
)

# ---- 5.1 パラメータ化 ----
add_heading(doc, '5.1 パラメータ化', level=2)

p = doc.add_paragraph()
p.add_run(
    'x 座標の代わりに正規化パラメータ t を使用する。'
)

add_formula(doc, 't = s / xSteps ∈ [0, 1]')
add_formula(doc, 's = 0, 1, 2, ..., xSteps  （x 方向のサンプル番号）')

p = doc.add_paragraph()
p.add_run('正規化の利点：')

items = [
    't = 0 がアゴ（固定端点）、t = 1 が切先（固定端点）に対応し、端点拘束を自然に表現できる',
    'アゴ〜切先の距離（ピクセル数）に依存しない汎用的な定式化になる',
    '数値計算上、係数行列の条件数が改善される（後述）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ---- 5.2 モデル式 ----
add_heading(doc, '5.2 曲線モデル式', level=2)

p = doc.add_paragraph()
p.add_run('求める曲線 y(t) を次式でモデル化する：')

add_formula(doc, 'y(t) = base(t) + t·(1 - t)·P(t)')

p = doc.add_paragraph()
p.add_run('各項の意味：')

# Sub-table for terms
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '記号'
hdr[1].text = '式'
hdr[2].text = '意味'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('base(t)', 'y₀ + (y₁ - y₀)·t', 'アゴとの切先を結ぶ弦（直線）'),
    ('t·(1 - t)', '（係数なし）', 'アゴ(t=0)と切先(t=1)でゼロになる「窓関数」'),
    ('P(t)', 'c₀ + c₁t + c₂t² + c₃t³', '刃線形状の偏差を表す多項式（DEG=3）'),
]
for row_data in data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        row.cells[i].text = text

doc.add_paragraph()

# ---- 5.3 端点拘束の保証 ----
add_heading(doc, '5.3 端点拘束の数学的保証', level=2)

p = doc.add_paragraph()
p.add_run(
    'モデル式が必ずアゴと切先を通ることを数学的に確認する。'
)

add_formula(doc, 't = 0 のとき：y(0) = base(0) + 0·(1-0)·P(0) = y₀ + 0 = y₀  ✓')
add_formula(doc, 't = 1 のとき：y(1) = base(1) + 1·(1-1)·P(1) = y₁ + 0 = y₁  ✓')

p = doc.add_paragraph()
p.add_run(
    'この保証は P(t) の値に依存しない。つまり、最小二乗法で係数 c₀〜c₃ がどのような値に'
    '決まったとしても、端点通過は常に満たされる。これが「端点拘束付き」と呼ぶ理由である。'
)

p = doc.add_paragraph()
p.add_run(
    '（参考）端点拘束なしの従来アプローチとの比較：単純に y(t) = a₀ + a₁t + ... + aₙtⁿ と'
    '置いて最小二乗フィットを行うと、係数が端点に引き寄せられる保証がなく、'
    'アゴ・切先からわずかに外れた曲線が得られる場合がある。'
)

# ---- 5.4 基底関数 ----
add_heading(doc, '5.4 基底関数', level=2)

p = doc.add_paragraph()
p.add_run(
    '偏差 r(t) = yMed(t) − base(t) を以下の基底関数の線形結合で近似する：'
)

add_formula(doc, 'r(t) ≈ t·(1-t)·P(t) = Σₖ cₖ · φₖ(t),  k = 0, 1, ..., DEG')
add_formula(doc, 'φₖ(t) = t·(1-t)·tᵏ = tᵏ⁺¹·(1-t)')

p = doc.add_paragraph()
p.add_run('各基底関数の形状：')

data = [
    ('k=0', 'φ₀(t) = t(1-t)', '対称な上に凸の二次曲線（ベータ分布 B(2,2) に比例）'),
    ('k=1', 'φ₁(t) = t²(1-t)', '右寄りの非対称山形。切先側の曲率変化を表現'),
    ('k=2', 'φ₂(t) = t³(1-t)', 'さらに右寄り。切先付近の急峻な変化を捉える'),
    ('k=3', 'φ₃(t) = t⁴(1-t)', '切先直前の急激な反りを表現。DEG=3 で追加'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '基底'
hdr[1].text = '式'
hdr[2].text = '形状の特徴'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
for k, formula, desc in data:
    row = table.add_row()
    row.cells[0].text = k
    row.cells[1].text = formula
    row.cells[2].text = desc

doc.add_paragraph()

add_heading(doc, '5.5 基底関数の数値的安定性', level=2)

p = doc.add_paragraph()
p.add_run(
    'φₖ(t) = tᵏ⁺¹(1-t) は通常の単項式基底 {1, t, t², ...} と類似しているが、'
    't の区間 [0,1] において以下の利点がある。'
)

items = [
    '全基底が区間 [0,1] 内でピークを持ち、値域が [0, ~0.25] 程度に抑えられる（単項式の tⁿ が区間端で急峻になる問題を軽減）',
    '各基底がゼロピークの位置を（k/(k+2) 付近に）ずらして持つため、行列 A が対角優位に近い構造を持つ',
    'DEG=3 では行列サイズが 4×4 と小さく、倍精度浮動小数点演算での数値誤差は無視できる水準',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run(
    '（注意）DEG を大きくすると高次の基底 φₖ(t) = tᵏ⁺¹(1-t) は t≈0 付近で急激に小さくなり、'
    '単項式基底の場合と同様に係数行列の条件数が悪化する可能性がある。'
    'DEG ≤ 5 程度が実用的な上限と考えられる。'
)

# ---- 5.6 正規方程式 ----
add_heading(doc, '5.6 最小二乗正規方程式', level=2)

p = doc.add_paragraph()
p.add_run(
    '残差の二乗和 Σ(r(tₛ) - Σcₖφₖ(tₛ))² を最小化するための正規方程式は：'
)

add_formula(doc, 'A · c = b')
add_formula(doc, 'Aₖⱼ = Σₛ φₖ(tₛ) · φⱼ(tₛ)')
add_formula(doc, 'bₖ   = Σₛ φₖ(tₛ) · r(tₛ)')
add_formula(doc, 'r(tₛ) = yMed[s] − base(tₛ)  （弦からの偏差）')

p = doc.add_paragraph()
p.add_run('s は s = 0, 1, ..., xSteps のすべてのサンプル点にわたる和を表す。')

add_heading(doc, '5.7 滑らかさの理由：自由度と過決定系', level=2)

p = doc.add_paragraph()
p.add_run(
    '正規方程式の意味を直感的に理解するための比較を以下に示す。'
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '項目'
hdr[1].text = '移動平均（window=3%）'
hdr[2].text = '多項式フィット（DEG=3）'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('自由度', '各点が独立に動く（高自由度）', '係数 4 個のみ（低自由度）'),
    ('データ点数', '約 1800 点', '約 1800 点（同一）'),
    ('局所ノイズの影響', '窓内のノイズが局所的に残る', '全体の平均を通じてノイズが相殺'),
    ('大域形状', '滑らかだが波打ちが残る場合あり', '大域的に最適な一本の多項式曲線'),
    ('切先付近の鋭い反り', '比較的追従しやすい', 'DEG=3 で十分表現可能'),
]
for row_data in data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        row.cells[i].text = text

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run(
    '約 1800 点のデータを 4 係数で近似する「過決定系（over-determined system）」では、'
    '最小二乗解が全データを最も良く説明するグローバル最適解となる。'
    '局所的なノイズは全体の最小化の中で自然に平均化される。'
)

# ---- 5.8 次数 DEG=3 の根拠 ----
add_heading(doc, '5.8 次数 DEG = 3 の根拠', level=2)

p = doc.add_paragraph()
p.add_run(
    '次数 DEG の選択は精度と過学習（overfitting）のトレードオフである。'
)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'DEG'
hdr[1].text = '全体次数'
hdr[2].text = '特徴'
hdr[3].text = '評価'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('1', '3 次', '対称放物線に近い。非対称刃形状を表現できない', '不足'),
    ('2', '4 次', 'ある程度の非対称性を表現できる', '許容'),
    ('3', '5 次', '切先付近の急峻な反りを自然に表現できる', '採用'),
    ('4 以上', '6 次以上', '過学習のリスク。端点付近で振動が生じる場合がある', '過剰'),
]
for row_data in data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        row.cells[i].text = text

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('合成データによる検証結果：')

add_formula(doc, '真の曲線 yTrue[s] + ±8px の一様ランダムノイズ → フィット後の平均誤差 ≈ 0.08px')

p = doc.add_paragraph()
p.add_run(
    'DEG=2 でも平均誤差は同程度だが、実際の包丁画像では切先付近の形状が非対称であるため、'
    'DEG=3 がより自然な曲線を与える。DEG=4 以上では合成データでの精度向上が見られず、'
    '実画像では過学習のリスクが高まる。'
)

# ---- 5.9 ガウス消去法 ----
add_heading(doc, '5.9 正規方程式の求解：部分ピボット選択付きガウス消去法', level=2)

p = doc.add_paragraph()
p.add_run(
    '正規方程式 A·c = b は 4×4 の連立一次方程式であり、'
    '部分ピボット選択付きガウス消去法 (Gaussian elimination with partial pivoting) で解く。'
)

add_heading(doc, 'アルゴリズムの概要', level=3)

items = [
    '前進消去（Forward elimination）：各列 k について、絶対値最大の行をピボットとして選択し行交換。上三角行列を形成',
    '後退代入（Back substitution）：上三角行列から未知数を上から順に求解',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

p = doc.add_paragraph()
p.add_run('実装コード：')

add_code_block(doc,
'function solveLinear(A, b) {\n'
'  const n = b.length;\n'
'  // 拡大係数行列 [A|b] を構築\n'
'  const M = A.map((row, i) => [...row, b[i]]);\n'
'  for (let col = 0; col < n; col++) {\n'
'    // 部分ピボット選択\n'
'    let maxRow = col;\n'
'    for (let r = col + 1; r < n; r++)\n'
'      if (Math.abs(M[r][col]) > Math.abs(M[maxRow][col])) maxRow = r;\n'
'    [M[col], M[maxRow]] = [M[maxRow], M[col]];\n'
'    if (Math.abs(M[col][col]) < 1e-12) return null;  // 特異行列\n'
'    // 前進消去\n'
'    for (let r = col + 1; r < n; r++) {\n'
'      const f = M[r][col] / M[col][col];\n'
'      for (let c = col; c <= n; c++) M[r][c] -= f * M[col][c];\n'
'    }\n'
'  }\n'
'  // 後退代入\n'
'  const x = new Array(n).fill(0);\n'
'  for (let i = n - 1; i >= 0; i--) {\n'
'    x[i] = M[i][n];\n'
'    for (let j = i + 1; j < n; j++) x[i] -= M[i][j] * x[j];\n'
'    x[i] /= M[i][i];\n'
'  }\n'
'  return x;\n'
'}'
)

add_heading(doc, '部分ピボット選択の必要性', level=3)

p = doc.add_paragraph()
p.add_run(
    '基底関数 φₖ(t) はすべて t(1-t) を因子に持つため、t≈0 や t≈1 付近（アゴや切先の近傍）では'
    '値が非常に小さくなる。これにより係数行列 A の対角成分が小さい行が生じる可能性があり、'
    'ピボットなしのガウス消去法では数値誤差が増幅される。'
    '部分ピボット選択により絶対値最大の要素を選ぶことで、数値安定性が確保される。'
)

p = doc.add_paragraph()
p.add_run(
    '（実際には DEG=3, n=4 の小規模行列では条件数が問題になることはほとんどないが、'
    'DEG を大きくした場合の安全性のためにピボット選択を実装している。）'
)

# ---- 5.10 フォールバック ----
add_heading(doc, '5.10 フォールバック処理', level=2)

p = doc.add_paragraph()
p.add_run(
    '正規方程式の係数行列 A がほぼ特異（行列式≈0）の場合、'
    'solveLinear は null を返す。この状況は以下のケースで生じうる：'
)

items = [
    'エッジ画像がほぼ真っ白または真っ黒で、すべての r(tₛ) がゼロに近い',
    'アゴと切先が同一点または極めて近い（xSteps が 0 か非常に小さい）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run(
    'フォールバック処理として、多項式フィットをスキップし、'
    'メディアンフィルタ後のデータ yMed をそのまま曲線データとして使用する。'
    '「フォールバック：多項式フィット失敗、メディアン値を直接使用」とコンソールに記録される。'
)

# ---- 5.11 スプライン代替案 ----
add_heading(doc, '5.11 スプライン代替案との比較', level=2)

p = doc.add_paragraph()
p.add_run(
    '三次スプラインは区間ごとに三次多項式を繋げた曲線で、より複雑な形状を表現できるが、'
    '本アプリでは多項式フィットを選択した。その理由を以下に示す。'
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '観点'
hdr[1].text = '三次スプライン'
hdr[2].text = '多項式フィット（本手法）'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('実装複雑さ', '比較的複雑（連続条件の連立方程式）', '単純（正規方程式 4×4 のみ）'),
    ('制御点数', 'ノット数に応じて増加', '係数 4 個のみ（固定）'),
    ('端点拘束', '別途条件を設定する必要あり', 'モデル式に組み込まれている'),
    ('振動リスク', 'Runge 現象は少ないが過学習あり', '低次数のため振動が起きにくい'),
    ('包丁形状への適合', '多くの制御点で細部まで追従', '全体的な滑らかさを優先'),
]
for row_data in data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        row.cells[i].text = text

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run(
    'ロボット研磨ルート生成においては、刃線の局所的な細部よりも「大域的に滑らかな曲線」が'
    '求められるため、4 係数の多項式フィットが適切な選択である。'
    '将来的に刃先の微細な形状（鎬筋の曲率変化など）を精密に追うニーズが生じた場合は'
    'スプラインへの変更を検討する余地がある。'
)

# ============================================================
# 6. 全体フローまとめ
# ============================================================
add_heading(doc, '6. 全体処理フロー', level=1)

add_code_block(doc,
'入力：エッジ検出画像、アゴ座標 (x₀,y₀)、切先座標 (x₁,y₁)\n'
'\n'
'ステップ①  エッジ追跡\n'
'  for s = 0 to xSteps:\n'
'    x   = x₀ + signX · s\n'
'    refY = round(y₀ + (y₁-y₀) · s/xSteps)  ← 弦上の参照点\n'
'    yRaw[s] = 下から上へ走査して最初に見つかったエッジ y\n'
'  yRaw[0] = y₀, yRaw[xSteps] = y₁  （端点固定）\n'
'\n'
'ステップ②  ギャップ補間\n'
'  null の連続区間を隣接検出点間の線形補間で埋める\n'
'\n'
'ステップ③  メディアンフィルタ\n'
'  window = max(3, round(xSteps × 0.01))\n'
'  yMed[s] = median(yRaw[s-w .. s+w])\n'
'\n'
'ステップ④  多項式フィット\n'
'  for s = 0 to xSteps:\n'
'    t    = s / xSteps\n'
'    base = y₀ + (y₁-y₀)·t\n'
'    r    = yMed[s] - base\n'
'    φₖ  = t(1-t)·tᵏ  (k=0..3)\n'
'    A[k][j] += φₖ·φⱼ  （正規行列を蓄積）\n'
'    b[k]    += φₖ·r    （右辺ベクトルを蓄積）\n'
'\n'
'  c = solveLinear(A, b)  ← 部分ピボット付きガウス消去法\n'
'\n'
'  for s = 0 to xSteps:\n'
'    t = s / xSteps\n'
'    P = c₀ + c₁t + c₂t² + c₃t³\n'
'    yCurve[s] = round(y₀ + (y₁-y₀)·t + t(1-t)·P)\n'
'\n'
'出力：滑らか曲線点列 {(x, yCurve[s])} および曲線長（弧長積分値）'
)

doc.add_paragraph()

# ============================================================
# 7. 曲線長の計算
# ============================================================
add_heading(doc, '7. 曲線長（弧長）の計算', level=1)

p = doc.add_paragraph()
p.add_run(
    'ロボット研磨ルートの経路長推定のため、求めた曲線の弧長 L を計算する。'
)

add_formula(doc, 'L = Σₛ √((Δx)² + (Δy)²)')
add_formula(doc, 'Δx = x[s] - x[s-1],  Δy = yCurve[s] - yCurve[s-1]')

p = doc.add_paragraph()
p.add_run(
    'このピクセル単位の弧長を、カード校正から求めたスケール係数（mm/px）で'
    'ミリメートル単位に変換した値が「曲線長 (mm)」として計測結果に表示される。'
)

doc.add_paragraph()

# ============================================================
# 8. まとめ
# ============================================================
add_heading(doc, '8. まとめ', level=1)

p = doc.add_paragraph(
    '本手法の特長を以下にまとめる。'
)

items = [
    '端点拘束：モデル式に t(1-t) を組み込むことで、アゴ・切先を必ず通ることを数学的に保証',
    '大域的滑らかさ：約 1800 点を 4 係数で近似する過決定系最小二乗法により、局所ノイズが自然に平均化される',
    '非対称形状への対応：DEG=3（全体 5 次）の多項式が切先付近の急峻な曲率変化を自然に表現',
    'ロバスト性：メディアンフィルタが単発のエッジ誤検知を除去し、多項式フィットが残差ノイズを平均化',
    '実装の簡潔さ：4×4 の線形方程式を Gaussian elimination で解くだけで実現。外部ライブラリ不要',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# Save
# ============================================================
out_path = '/home/user/2026/blade_curve_method.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
